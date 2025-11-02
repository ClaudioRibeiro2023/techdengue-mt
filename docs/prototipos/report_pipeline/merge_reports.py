#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
merge_reports.py
Preenche templates DOCX (EPI01, EVD01, OP01) com dados JSON e, se disponíveis,
converte para PDF e PDF/A-1.
Requisitos: python-docx, (opcional) LibreOffice para PDF, Ghostscript para PDF/A-1.
"""
import os
import sys
import json
import argparse
import shutil
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def replace_in_paragraph(paragraph, mapping):
    """Substitui placeholders simples {{chave}} no parágrafo."""
    if not paragraph.text:
        return
    for k, v in mapping.items():
        token = "{{" + k + "}}"
        if token in paragraph.text:
            paragraph.text = paragraph.text.replace(token, str(v))

def replace_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            replace_in_cell(cell, mapping)

def replace_in_cell(cell, mapping):
    for p in cell.paragraphs:
        replace_in_paragraph(p, mapping)
    for t in cell.tables:
        replace_in_table(t, mapping)

def replace_placeholders(doc, mapping):
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for t in doc.tables:
        replace_in_table(t, mapping)

def insert_images(doc, image_map):
    """
    Insere imagens onde encontrar parágrafos contendo tokens do tipo:
    [IMG:key:width_in_inches:align]
    Exemplo: [IMG:chart_tendencia:5:CENTER]
    O caminho da imagem virá de image_map[key]
    """
    to_insert = []
    for i, p in enumerate(doc.paragraphs):
        if p.text and p.text.strip().startswith("[IMG:") and p.text.strip().endswith("]"):
            content = p.text.strip()[5:-1]  # remove [IMG: e ]
            parts = content.split(":")
            if len(parts) >= 1:
                key = parts[0]
                width = float(parts[1]) if len(parts) >= 2 else 5.0
                align = parts[2].upper() if len(parts) >= 3 else "LEFT"
                to_insert.append((i, key, width, align))

    # Inserir de trás pra frente para não deslocar índices
    for i, key, width, align in reversed(to_insert):
        par = doc.paragraphs[i]
        par.clear()
        run = par.add_run()
        img_path = image_map.get(key)
        if not img_path or not os.path.exists(img_path):
            par.add_run(f"[imagem não encontrada: {key}]").bold = True
            continue
        run.add_picture(img_path, width=Inches(width))
        if align == "CENTER":
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "RIGHT":
            par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT

def convert_to_pdf(docx_path, out_dir):
    """Tenta converter via LibreOffice (soffice). Retorna caminho do PDF ou None."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    cmd = [soffice, "--headless", "--convert-to", "pdf", docx_path, "--outdir", out_dir]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        pdf_path = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception:
        return None

def convert_to_pdfa(pdf_path, out_dir):
    """Tenta converter via Ghostscript para PDF/A-1b. Retorna caminho do PDF/A ou None."""
    gs = shutil.which("gs") or shutil.which("gswin64c") or shutil.which("gswin32c")
    if not gs:
        return None
    out_pdfa = os.path.join(out_dir, os.path.splitext(os.path.basename(pdf_path))[0] + "_PDFA.pdf")
    cmd = [
        gs, "-dPDFA=1", "-dBATCH", "-dNOPAUSE",
        "-sProcessColorModel=DeviceRGB",
        "-sDEVICE=pdfwrite",
        f"-sOutputFile={out_pdfa}",
        "-sPDFACompatibilityPolicy=1",
        pdf_path
    ]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return out_pdfa if os.path.exists(out_pdfa) else None
    except Exception:
        return None

def main():
    ap = argparse.ArgumentParser(description="Merge de relatórios DOCX → PDF → PDF/A-1")
    ap.add_argument("--template", required=True, help="Caminho do template DOCX")
    ap.add_argument("--data", required=True, help="JSON com dados (placeholders)")
    ap.add_argument("--images", help="JSON {key: caminho_imagem}")
    ap.add_argument("--out", default="out.docx", help="Arquivo DOCX de saída")
    ap.add_argument("--pdf", action="store_true", help="Também gerar PDF (LibreOffice)")
    ap.add_argument("--pdfa", action="store_true", help="Também gerar PDF/A-1 (Ghostscript)")
    args = ap.parse_args()

    # Carrega dados e imagens
    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)
    image_map = {}
    if args.images and os.path.exists(args.images):
        with open(args.images, "r", encoding="utf-8") as f:
            image_map = json.load(f)

    # Abre template e executa merge simples
    doc = Document(args.template)
    replace_placeholders(doc, data)
    insert_images(doc, image_map)

    # Salva DOCX
    out_docx = args.out
    os.makedirs(os.path.dirname(out_docx) or ".", exist_ok=True)
    doc.save(out_docx)
    print(f"[OK] DOCX gerado: {out_docx}")

    # Opções de PDF e PDF/A-1
    if args.pdf or args.pdfa:
        out_dir = os.path.dirname(out_docx) or "."
        pdf_path = convert_to_pdf(out_docx, out_dir) if args.pdf or args.pdfa else None
        if pdf_path:
            print(f"[OK] PDF gerado: {pdf_path}")
            if args.pdfa:
                pdfa_path = convert_to_pdfa(pdf_path, out_dir)
                if pdfa_path:
                    print(f"[OK] PDF/A-1 gerado: {pdfa_path}")
                else:
                    print("[WARN] Ghostscript não encontrado ou falha na conversão para PDF/A-1.")
        else:
            print("[WARN] LibreOffice (soffice) não encontrado ou falha na conversão para PDF.")

if __name__ == "__main__":
    main()
